#!/usr/bin/env python3
"""Generate `Why-GateOverflow-Chatbot.docx` — a positioning document explaining
why this tool helps GATE CS & DA preparation better than general AI chatbots,
plus a roadmap of differentiating features.

Run:  python demo/build_doc.py   (needs python-docx: pip install python-docx)
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Why-GateOverflow-Chatbot.docx"

RED = RGBColor(0xDC, 0x26, 0x26)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x55, 0x5B, 0x66)


def main() -> int:
    doc = Document()

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def heading(text: str, level: int = 1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RED if level == 1 else DARK
        return h

    def para(text: str = "", *, bold=False, italic=False, size=11, color=None,
             align=None, space_after=6):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def bullet(lead: str, rest: str = "", level: int = 0):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Pt(18 + level * 16)
        r = p.add_run(lead)
        r.bold = True
        if rest:
            p.add_run(" — " + rest if not rest.startswith(" ") else rest)
        return p

    # ---------------------------------------------------------------- Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("GateOverflow Chatbot")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RED
    para("Why a GATE-specialist AI assistant beats general chatbots for "
         "CS & DA exam preparation", italic=True, size=12, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para("Your AI study buddy for GATE CS & DA — grounded in real previous-year "
         "questions.", italic=True, size=10, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    # ---------------------------------------------------------------- Intro
    heading("The question", level=1)
    para("Tools like ChatGPT, Claude and DeepSeek are powerful generalists. So "
         "why use a dedicated GateOverflow Chatbot for GATE Computer Science (CS) "
         "and Data Science & AI (DA) preparation? Because exam prep needs "
         "something general chatbots are not built for: answers that are "
         "trustworthy, exam-aligned, grounded in real previous-year questions, "
         "and able to learn from your study material and feedback.")

    # ---------------------------------------------------- Core difference
    heading("The core difference: grounded, not guessing", level=1)
    bullet("Answers from real GATE material, not generic web knowledge",
           "it retrieves from your ingested previous-year questions (PYQs) and "
           "notes — GATE CS/DA, ISRO, NIELIT, UGC-NET, TIFR — so answers match the "
           "exam's style, depth and expected method instead of wandering into a "
           "generic textbook tangent.")
    bullet("Every answer shows its sources",
           "file name, page and a relevance score. You can trust and verify "
           "rather than wondering whether it was hallucinated — critical when one "
           "wrong formula costs you marks.")
    bullet("No hallucinated exam facts",
           "it is instructed to ground answers in retrieved material and to flag "
           "when it is relying on general knowledge, reducing the confident-but-"
           "wrong answers that hurt you on objective exams.")

    # ---------------------------------------------------- Built for GATE
    heading("Built specifically for the GATE workflow", level=1)
    bullet("PYQ-first practice",
           "\"give me a PYQ on TLB and solve it\" pulls actual previous-year "
           "questions and walks the method — general chatbots often invent "
           "plausible-but-fake \"GATE questions\".")
    bullet("Exam-tuned answers",
           "step-by-step numericals, complexity bounds, proper LaTeX, and the "
           "\"why each option is right/wrong\" reasoning GATE rewards.")
    bullet("Interactive quizzing",
           "\"quiz me with 5 MCQs, one at a time, then score me\" turns prep into "
           "active recall — the most effective study method.")
    bullet("Stays on syllabus",
           "a focused scope means fewer distractions; it nudges you back to "
           "CS/DS/ML/AI instead of wandering off-topic.")

    # ---------------------------------------------------- Daily advantages
    heading("Practical daily-use advantages", level=1)
    bullet("Bring your own material",
           "drop your handwritten notes, a coaching PDF or a textbook chapter and "
           "it answers from that — so it speaks your syllabus and your sources.")
    bullet("Attach a doubt as a photo or PDF",
           "snap a question you're stuck on; the vision model reads and solves it "
           "— no retyping.")
    bullet("It improves from your feedback (RLHF)",
           "thumbs-down a weak answer, say why, and it regenerates a better one; "
           "your corrections build a preference dataset to fine-tune the model "
           "over time. General chatbots forget your corrections.")
    bullet("Cost-efficient and privacy-friendly",
           "runs on a cheap model by default, your study data stays in your own "
           "storage and database, and you control what it learns from.")
    bullet("A focused, distraction-free study surface",
           "conversation memory, follow-up suggestions, copy/regenerate — a study "
           "tool, not a general assistant you must constantly steer toward GATE.")

    # ---------------------------------------------------- Extra edges
    heading("More reasons it stands apart from market AI tools", level=1)
    bullet("Domain-curated knowledge base",
           "seeded with dense CS/DS/ML/AI syllabus notes and a one-command fetcher "
           "for free, openly-licensed textbooks — knowledge-rich out of the box.")
    bullet("Transparent and auditable",
           "citations + an \"answered from general knowledge\" notice tell you how "
           "confident/grounded each answer is.")
    bullet("Owns its data and learning loop",
           "conversations, feedback and preference pairs are stored and exportable "
           "as JSONL for reward-model / DPO training — a continuously improving, "
           "self-hostable system rather than a black box.")
    bullet("Built to production standards",
           "React + TypeScript + Tailwind front-end, FastAPI backend, Dockerised, "
           "tested and configurable — deployable for a whole cohort, not just a "
           "personal chat window.")
    bullet("On-brand and integrable",
           "carries the GateOverflow identity and is designed to plug into "
           "GateOverflow's existing PYQ bank, user accounts and rank data.")
    bullet("Model-flexible",
           "point CHAT_MODEL at a stronger model when you want deeper reasoning — "
           "you get specialist grounding and frontier reasoning together.")

    # ---------------------------------------------------- Summary
    heading("In one line", level=1)
    p = doc.add_paragraph()
    p.add_run("ChatGPT and Claude are brilliant generalists that can "
              "hallucinate; the GateOverflow Chatbot is a GATE specialist that "
              "cites real previous-year questions and your own material, "
              "practises you with PYQs and quizzes, learns from your feedback, "
              "and stays on-syllabus — exactly what daily exam prep needs.").italic = True

    para("Honest caveat: for raw reasoning on a brand-new, never-seen problem a "
         "frontier model may go deeper. This tool's edge is trust, exam-alignment, "
         "your-material grounding and PYQ practice — and you can always point it at "
         "a stronger model to get both.", color=GREY, size=10, space_after=10)

    # ---------------------------------------------------- Roadmap
    heading("Features that would make it even more differentiated", level=1)
    para("A roadmap of capabilities that general chatbots cannot easily match, "
         "because they depend on GATE-specific data, persistence and workflow:",
         space_after=8)

    def sub(t):
        heading(t, level=2)

    sub("Personalization & adaptive learning")
    bullet("Spaced-repetition engine",
           "tracks weak topics and re-quizzes them on an optimal schedule (SM-2).")
    bullet("Adaptive difficulty",
           "questions get harder/easier based on your recent accuracy.")
    bullet("Personalized study planner",
           "a day-by-day plan to your exam date, rebalanced as you progress.")

    sub("Assessment & analytics")
    bullet("Full mock-test simulator",
           "a timed 3-hour CBT with MCQ/MSQ/NAT, negative marking and an instant "
           "score + estimated GATE rank/percentile.")
    bullet("Performance dashboard",
           "accuracy by subject, time-per-question, mastery heatmap and "
           "exam-readiness score.")
    bullet("Weak-area detection",
           "automatically surfaces the topics costing you the most marks.")

    sub("Retrieval & accuracy")
    bullet("Hybrid search + reranking",
           "BM25 + dense retrieval with a cross-encoder reranker for sharper, "
           "more relevant citations.")
    bullet("Citation click-through",
           "jump straight to the exact source page/PDF, or an NPTEL video "
           "timestamp.")
    bullet("Confidence & self-check",
           "a verifier pass that flags low-confidence or unsupported claims.")

    sub("Multimodal & accessibility")
    bullet("Voice mode",
           "ask by speaking and hear answers read aloud — great for revision on "
           "the go.")
    bullet("Handwriting & diagram input",
           "photograph your worked solution and get step-by-step feedback.")
    bullet("Multilingual explanations",
           "Hindi and regional-language explanations alongside English.")
    bullet("Socratic tutor mode",
           "gives hints first and reveals the full solution only when you're "
           "stuck — builds real understanding.")

    sub("Engagement & community")
    bullet("Daily question, streaks & gamification",
           "habit-forming practice with badges and leaderboards.")
    bullet("Auto-generated flashcards & formula sheets",
           "turn any chat into revision cards and a printable formula sheet.")
    bullet("Community integration",
           "tie into GateOverflow's real Q&A threads, upvotes and discussions for "
           "human + AI answers side by side.")

    sub("Platform & integration")
    bullet("GateOverflow account & profile sync",
           "personalize using a user's real activity, attempted questions and "
           "rank history.")
    bullet("Email/calendar digests",
           "weekly weak-area summaries and revision reminders.")
    bullet("Local/offline model option",
           "a privacy-first, low-cost on-device mode for sensitive notes.")

    doc.add_paragraph()
    para("Document generated for the GateOverflow Chatbot demo kit · "
         "18.Utilities/E.GATE-CS-DA-Chatbot", italic=True, size=9, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUT)
    print(f"✅ Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
