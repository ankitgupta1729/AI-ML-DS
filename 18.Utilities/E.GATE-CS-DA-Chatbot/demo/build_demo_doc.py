#!/usr/bin/env python3
"""Generate `GateOverflow-Chatbot-Demo.docx` — the complete team-demo document:
features, functionalities, workflow, architecture, differentiation vs general
chatbots, sample Q&A, and a roadmap of improvements.

Run:  python demo/build_demo_doc.py   (needs: pip install python-docx)
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent / "GateOverflow-Chatbot-Demo.docx"
RED = RGBColor(0xDC, 0x26, 0x26)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x55, 0x5B, 0x66)
GREEN = RGBColor(0x16, 0xA3, 0x4A)


def main() -> int:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ---- helpers -------------------------------------------------------- #
    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.color.rgb = RED if level == 1 else DARK
        return h

    def para(text="", *, bold=False, italic=False, size=10.5, color=None, align=None, after=6):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold, r.italic, r.font.size = bold, italic, Pt(size)
        if color is not None:
            r.font.color.rgb = color
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        return p

    def bullet(lead, rest="", level=0):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Pt(18 + level * 16)
        p.add_run(lead).bold = True
        if rest:
            p.add_run(" — " + rest)
        return p

    def numbered(text):
        return doc.add_paragraph(text, style="List Number")

    def shade(cell, hex_):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), hex_)
        tcPr.append(sh)

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            shade(c, "DC2626")
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(val)
                run.font.size = Pt(9.5)
        return t

    # ---- Title ---------------------------------------------------------- #
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("GateOverflow Chatbot")
    tr.bold = True
    tr.font.size = Pt(30)
    tr.font.color.rgb = RED
    para("Complete Demo & Documentation — Features, Workflow, Architecture & Differentiation",
         italic=True, size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para("Your AI study buddy for GATE CS & DA — grounded in real previous-year questions.",
         italic=True, size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

    # ---- 1. Executive summary ------------------------------------------ #
    heading("1. Executive summary", 1)
    para("The GateOverflow Chatbot is a production-ready, Retrieval-Augmented "
         "Generation (RAG) study platform for the GATE Computer Science (CS) and "
         "Data Science & AI (DA) exams, built for GateOverflow (gateoverflow.in). "
         "It is far more than a chat box: it is a full preparation workspace that "
         "answers questions grounded in real material with citations, generates "
         "and scores mock tests with GATE rules, schedules revision with spaced "
         "repetition, plans study time, tracks performance, and continuously "
         "improves from human feedback.")
    para("This document covers everything for a team demo: the full feature set, "
         "how it works end-to-end, the architecture, how it differs from general "
         "chatbots, sample interactions, and a roadmap of refinements.", color=GREY)

    # ---- 2. Differentiation -------------------------------------------- #
    heading("2. Why it's different from ChatGPT, Claude & DeepSeek", 1)
    para("General chatbots are excellent generalists — but exam preparation needs "
         "trust, exam-alignment and structure. The table below summarises the gap.")
    table(
        ["Capability", "ChatGPT / Claude / DeepSeek", "GateOverflow Chatbot"],
        [
            ["Source of answers", "Generic web-scale knowledge", "Your PYQs, notes & curated material (RAG)"],
            ["Trust / verifiability", "No citations; can hallucinate", "Citations: file, page & match score"],
            ["Previous-year questions", "Can invent fake 'GATE' questions", "Ingested real PYQs (CS, DA, ISRO, NIELIT, UGC-NET, TIFR)"],
            ["Mock tests & scoring", "None", "MCQ/MSQ/NAT, GATE negative marking, percentile"],
            ["Revision", "None", "Spaced repetition (SM-2) + auto flashcards"],
            ["Progress insight", "None", "Dashboard: readiness, weak areas, streaks"],
            ["Learns from feedback", "Forgets corrections", "RLHF: preference pairs exported for DPO"],
            ["Your own material", "Cannot use it natively", "Ingest your PDFs/notes; answers from them"],
            ["Scope & focus", "Open-ended; easy to drift", "GATE-tuned; on-syllabus by design"],
            ["Deployment & data", "Hosted black box", "Self-hostable; your data, your DB"],
        ],
    )
    para("In one line: ChatGPT/Claude/DeepSeek are brilliant generalists that can "
         "hallucinate; the GateOverflow Chatbot is a GATE specialist that cites "
         "real questions and your material, practises and scores you, schedules "
         "revision, and learns from your feedback. You can even point it at a "
         "stronger model to get specialist grounding plus frontier reasoning.",
         italic=True, after=8)

    # ---- 3. Feature catalogue ------------------------------------------ #
    heading("3. Full feature catalogue", 1)

    heading("3.1 Conversational AI & RAG", 2)
    bullet("Grounded answers", "retrieval over ingested PDFs/notes; prefers your material over generic knowledge.")
    bullet("Citations & confidence", "source file, page/slide, subject and a similarity score; a grounding-confidence meter per answer.")
    bullet("Streaming UX", "token-by-token responses (Server-Sent Events).")
    bullet("Conversation memory", "history-aware follow-up rewriting (e.g. 'what about its complexity?').")
    bullet("Markdown, LaTeX & code", "KaTeX maths (auto-normalises \\(…\\)/\\[…\\] to $…$) and syntax-highlighted code.")
    bullet("Greeting & scope handling", "friendly small-talk; polite decline only for clearly off-topic asks.")

    heading("3.2 PYQ practice & mock tests", 2)
    bullet("Question generation", "exam-style MCQ, MSQ and NAT items, grounded in real PYQs.")
    bullet("Server-side scoring", "GATE negative marking (−1/3 for 1-mark, −2/3 for 2-mark MCQ; none for MSQ/NAT).")
    bullet("Percentile estimate", "from the score ratio, plus a timer and per-question explanations.")
    bullet("Closed-book integrity", "answers are withheld from the client and scored on the server.")

    heading("3.3 Spaced repetition & flashcards", 2)
    bullet("SM-2 scheduling", "the proven SuperMemo-2 algorithm schedules reviews by recall quality.")
    bullet("Auto-generated cards", "every wrong quiz answer becomes a flashcard automatically.")
    bullet("Topic generation", "generate a deck on any topic on demand; review with Again/Hard/Good/Easy.")

    heading("3.4 Study planner", 2)
    bullet("Day-by-day plan", "to your exam date, prioritising high-weight subjects with PYQs and revision.")
    bullet("Calendar export", "download the plan as an .ics file for any calendar app (offline).")

    heading("3.5 Performance dashboard & analytics", 2)
    bullet("Exam-readiness score", "blends accuracy, practice volume and study habit (0–100).")
    bullet("Accuracy by subject", "with automatic weak-area detection to prioritise.")
    bullet("Trends", "average percentile, streaks, due-review count and recent attempts.")

    heading("3.6 Daily question & gamification", 2)
    bullet("Question of the day", "a fresh, hinted question to build a daily habit.")
    bullet("Streaks", "consecutive-day activity tracking to reinforce consistency.")

    heading("3.7 Multimodal & accessibility", 2)
    bullet("Image attachments", "snap/upload a question; the vision model reads and solves it.")
    bullet("PDF / text attachments", "extracted and used as grounding context.")
    bullet("Sketch input", "draw a diagram/working on a canvas and send it as an image.")
    bullet("Voice in & out", "speak your question (speech-to-text) and have answers read aloud (TTS).")
    bullet("Multilingual answers", "explanations in Hindi and other Indian languages.")

    heading("3.8 Tutoring modes", 2)
    bullet("Socratic mode", "gives hints and leading questions first; reveals the full solution when you're ready.")
    bullet("Per-answer actions", "copy, like, dislike, regenerate and read-aloud on every reply, plus follow-up chips.")

    heading("3.9 Feedback & RLHF", 2)
    bullet("Rich feedback", "thumbs up/down with reason, comment and an optional 'better answer'.")
    bullet("Feedback-steered regeneration", "a disliked answer is re-generated using your feedback (online loop).")
    bullet("Preference dataset", "corrections/regenerations stored as (prompt, chosen, rejected) pairs.")
    bullet("JSONL export", "for offline reward-model / DPO fine-tuning, then redeploy an improved model.")

    heading("3.10 GateOverflow links, bookmarks, history & export", 2)
    bullet("Direct GateOverflow PYQ links", "for previous-year-question answers ONLY, the exact GateOverflow question thread(s) are shown at the end of the answer — extracted from the URLs embedded in the PYQ PDFs (gateoverflow.in/<id>/<slug>). Open the official discussion & community answers in one click.")
    bullet("Bookmarks & chat history", "star any answer to save it; revisit past conversations and saved answers from the history panel.")
    bullet("Cheat-sheet & PDF export", "turn a conversation into a concise revision cheat-sheet, or export the whole chat — print or save as PDF.")
    bullet("Optional reranking", "cross-encoder reranking of citations via the lightweight FlashRank package (set RERANK=true).")

    heading("3.11 Persistence, branding & platform", 2)
    bullet("Database", "SQLAlchemy (SQLite by default, Postgres-ready): conversations, messages, feedback, quizzes, reviews, plans, activity, bookmarks.")
    bullet("GateOverflow identity", "the 'GO + red ?' wordmark and brand colours; light & dark themes; responsive UI.")
    bullet("Local/offline model", "point at any OpenAI-compatible endpoint (Ollama, LM Studio, vLLM) via OPENAI_BASE_URL.")

    # ---- 4. Workflow ---------------------------------------------------- #
    heading("4. How it works — end-to-end workflow", 1)
    heading("4.1 Knowledge ingestion (one-time / incremental)", 2)
    numbered("Drop study material and PYQ PDFs into data/.")
    numbered("Files are parsed, split into overlapping chunks, tagged with provenance, embedded (OpenAI) and stored in Chroma with content-hash IDs.")
    numbered("A manifest records each file so re-runs only embed new/changed files.")

    heading("4.2 Asking a question", 2)
    numbered("The user types/speaks/attaches a question (optionally Socratic mode or a language).")
    numbered("Follow-ups are rewritten into a standalone query using chat history.")
    numbered("Top-k relevant chunks are retrieved from Chroma; the best score drives the confidence meter.")
    numbered("The LLM streams a grounded, exam-oriented answer; sources are de-duplicated and shown.")
    numbered("The turn (and any rating) is persisted to the database.")

    heading("4.3 The improvement (RLHF) loop", 2)
    numbered("User dislikes an answer and gives a reason / better answer.")
    numbered("Regeneration is steered by that feedback (immediate improvement).")
    numbered("A (prompt, chosen, rejected) preference pair is stored.")
    numbered("Export JSONL → train a reward model / run DPO offline → redeploy a better CHAT_MODEL.")

    heading("4.4 Practice & revision journey", 2)
    numbered("Generate a quiz/mock → answer → server scores with GATE rules → see explanations, percentile and weak areas.")
    numbered("Wrong answers become flashcards, scheduled by SM-2.")
    numbered("Review due cards daily; the dashboard tracks readiness; the planner keeps you on schedule.")

    # ---- 5. Architecture ------------------------------------------------ #
    heading("5. Architecture & technology stack", 1)
    para("Frontend (React + TypeScript + Vite + Tailwind v4) ⇄ Backend (FastAPI) "
         "⇄ RAG engine (Chroma vector store + OpenAI) with a SQLAlchemy database "
         "for all state. Answers stream over Server-Sent Events.")
    table(
        ["Layer", "Technology", "Responsibility"],
        [
            ["Frontend", "React 19, TypeScript, Vite 6, Tailwind v4", "Chat + study-suite UI, streaming, voice/sketch, themes"],
            ["API", "FastAPI, Uvicorn", "Chat, quiz, review, plan, feedback, analytics endpoints"],
            ["RAG", "LangChain, Chroma, OpenAI embeddings", "Retrieval-augmented, cited generation"],
            ["Model", "gpt-4o-mini (configurable; vision)", "Answers, quizzes, flashcards, plans"],
            ["Database", "SQLAlchemy (SQLite / Postgres)", "Conversations, feedback, quizzes, SR, plans, streaks"],
            ["Ingestion", "pypdf, docx2txt, python-pptx", "Parse → chunk → embed (incremental)"],
            ["Deploy", "Docker, docker-compose, nginx", "Frontend (nginx) + API services"],
        ],
    )

    heading("5.1 Key API endpoints", 2)
    table(
        ["Endpoint", "Purpose"],
        [
            ["POST /chat, /chat/stream", "Answer (sync / streamed SSE) with sources & confidence"],
            ["POST /feedback", "Like/dislike (+ reason, comment, correction)"],
            ["POST /regenerate/stream", "Regenerate steered by feedback (RLHF)"],
            ["GET /export/preferences", "Download RLHF preference dataset (JSONL)"],
            ["POST /quiz/generate, /quiz/submit", "Generate & score a quiz/mock test"],
            ["POST /flashcards/generate", "Create spaced-repetition cards on a topic"],
            ["GET /review/due, POST /review/grade", "Due cards & SM-2 grading"],
            ["POST /plan/generate, GET /plan, /plan/calendar.ics", "Study plan & calendar export"],
            ["GET /daily, GET /analytics", "Daily question + streak; dashboard data"],
            ["GET /conversations, /conversations/{id}", "Chat history (list + load)"],
            ["POST /bookmark, GET /bookmarks", "Save / list bookmarked answers"],
            ["POST /cheatsheet", "Build a revision cheat-sheet from a chat"],
            ["GET /meta, /health, /admin/stats", "Branding, liveness, usage counters"],
        ],
    )

    # ---- 6. Sample interactions ---------------------------------------- #
    heading("6. Sample interactions", 1)

    def qa(q, a_lines):
        p = doc.add_paragraph()
        p.add_run("You: ").bold = True
        p.add_run(q)
        p.paragraph_format.space_after = Pt(2)
        for line in a_lines:
            ap = doc.add_paragraph(style="List Bullet")
            r = ap.add_run("GO Buddy: " if line is a_lines[0] else "")
            r.bold = True
            r.font.color.rgb = RED
            ap.add_run(line)
        doc.add_paragraph()

    qa("Explain the time complexity of merge sort and why.",
       ["Merge sort is Θ(n log n). It splits the array in half each level and merges in linear time.",
        "Recurrence T(n) = 2T(n/2) + O(n) → by the Master Theorem, Θ(n log n).",
        "Stable, with worst-case O(n log n) unlike quicksort.  [Source: cs_core_concepts.md · 91% match]"])
    qa("Solve: TLB hit ratio 0.9, TLB access 10 ns, memory access 100 ns — effective access time?",
       ["Hit (0.9): 10 + 100 = 110 ns.  Miss (0.1): 10 + 100 + 100 = 210 ns.",
        "EAT = 0.9·110 + 0.1·210 = 120 ns.  ✓  [Source: filter1_volume2.pdf · p.57 · 88% match]"])
    qa("Hi",
       ["Hi there! I'm GO Buddy, your study companion from GateOverflow. Ask me about a concept, a PYQ, or your prep strategy — what shall we start with?  (No junk retrieval on greetings.)"])
    qa("Tell me a joke about cricket.",
       ["I'm focused on GATE CS / DA preparation, so I can't help with that — but ask me anything about the syllabus, a tricky PYQ, or your prep strategy!  (Polite out-of-scope handling.)"])

    # ---- 7. Security & production -------------------------------------- #
    heading("7. Security & production notes", 1)
    bullet("Secrets", "OPENAI_API_KEY lives only in .env (git-ignored); never committed.")
    bullet("CORS", "open in dev for convenience; restrict to the front-end origin (or use the nginx same-origin proxy) before deploying.")
    bullet("Auth & rate limiting", "add an API gateway / auth layer in front of FastAPI for production.")
    bullet("Resilience", "persistence is best-effort — a database hiccup never breaks a chat.")
    bullet("Cost control", "cheap defaults (gpt-4o-mini, text-embedding-3-small); incremental ingestion avoids re-embedding.")
    bullet("Copyright", "only ingest material you are licensed to use.")
    bullet("Testing", "offline smoke tests cover scoring, SM-2, JSON extraction and DB flows; the frontend type-checks and builds.")

    # ---- 8. Running & demo kit ----------------------------------------- #
    heading("8. Running the app & demo kit", 1)
    para("Run everything with one command: make dev  (or bash scripts/dev.sh), "
         "then open http://localhost:5173. The API runs on :8000. Docker: "
         "docker compose up --build (frontend on :8080).")
    bullet("Demo video", "demo/GateOverflow-Chatbot-Demo.mp4 — narrated ~3-minute walkthrough of every feature.")
    bullet("Wireframe", "demo/wireframe.html — annotated screens & UX spec (open in browser / import to Figma).")
    bullet("Demo script", "demo/DEMO_SCRIPT.md — storyboard, narration and sample Q&A for a live demo.")

    # ---- 9. Roadmap ----------------------------------------------------- #
    heading("9. Roadmap — further improvements & refinements", 1)
    heading("9.1 Retrieval & accuracy", 2)
    bullet("Hybrid search + cross-encoder reranking", "BM25 + dense retrieval with a reranker for sharper citations.")
    bullet("Citation click-through", "open the exact source page/PDF, or an NPTEL video timestamp.")
    bullet("Answer self-verification", "a checker pass that flags low-confidence or unsupported claims.")

    heading("9.2 Personalisation & assessment", 2)
    bullet("Adaptive difficulty", "questions adjust to recent accuracy.")
    bullet("Full timed CBT simulator", "3-hour exam mode with section timing and a virtual calculator.")
    bullet("Predicted rank", "map performance to an estimated AIR using historical distributions.")

    heading("9.3 Engagement & platform", 2)
    bullet("GateOverflow account sync", "personalise with a user's real attempts, profile and rank history (direct question-thread links are already implemented; account sync needs GateOverflow's API/OAuth).")
    bullet("Email / push digests", "weekly weak-area summaries and revision reminders (needs SMTP/push setup).")
    bullet("Community integration", "show AI answers beside GateOverflow's human Q&A and upvotes.")
    bullet("Gamification", "badges, leaderboards and shareable progress.")

    heading("9.4 Operations", 2)
    bullet("Evaluation harness", "benchmark answer quality against a labelled PYQ set per release.")
    bullet("Observability & caching", "request tracing, usage dashboards and response caching for cost.")
    bullet("Multi-tenant deployment", "shared vector DB and per-user data isolation for a whole cohort.")

    doc.add_paragraph()
    para("Built for GateOverflow · 18.Utilities/E.GATE-CS-DA-Chatbot · "
         "Generated for the team demo.", italic=True, size=9, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUT)
    print(f"✅ Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
